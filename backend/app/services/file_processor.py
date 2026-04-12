# File processing service for different file types
import PyPDF2
import csv
import io
from docx import Document as DocxDocument
import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    """Processes different file types and extracts text"""

    @staticmethod
    def process_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF using PyPDF2.

        Args:
            file_content: Raw PDF bytes

        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            extracted_text = "\n".join(text_parts)
            logger.info(f"PDF processed: {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    @staticmethod
    def process_csv(file_content: bytes) -> str:
        """
        Extract text from CSV.

        Args:
            file_content: Raw CSV bytes

        Returns:
            Extracted text formatted as readable content
        """
        try:
            csv_file = io.StringIO(file_content.decode('utf-8'))
            csv_reader = csv.DictReader(csv_file)

            text_parts = ["CSV Data:\n"]
            for row_num, row in enumerate(csv_reader, 1):
                row_text = " | ".join([f"{k}: {v}" for k, v in row.items()])
                text_parts.append(f"Row {row_num}: {row_text}")

            extracted_text = "\n".join(text_parts)
            logger.info(f"CSV processed: {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            logger.error(f"CSV processing failed: {str(e)}")
            raise ValueError(f"Failed to process CSV: {str(e)}")

    @staticmethod
    def process_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_content: Raw DOCX bytes

        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_data))

            extracted_text = "\n".join(text_parts)
            logger.info(f"DOCX processed: {len(extracted_text)} characters")
            return extracted_text

        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    @staticmethod
    def process_file(file_content: bytes, file_type: str) -> str:
        """
        Route file to appropriate processor based on type.

        Args:
            file_content: Raw file bytes
            file_type: File extension (pdf, csv, docx)

        Returns:
            Extracted text
        """
        file_type = file_type.lower()

        if file_type == "pdf":
            return FileProcessor.process_pdf(file_content)
        elif file_type == "csv":
            return FileProcessor.process_csv(file_content)
        elif file_type == "docx":
            return FileProcessor.process_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
